"""
Chatty — Odoo PDF & attachment tools.

Download Odoo document PDFs via QWeb report engine, read/create/list
attachments on Odoo records. Text extraction for PDFs, Word docs,
spreadsheets, CSV, and plain text files. File cache system for
forwarding attachments by reference.
"""

import base64
import io
import logging
import xmlrpc.client

from ..constants import ALLOWED_MODELS
from ..helpers import safe_get_client

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB

# ── File-type classification ──────────────────────────────────────────

_TEXT_EXTENSIONS = {
    ".txt", ".eml", ".html", ".htm", ".xml", ".json", ".log",
    ".md", ".yaml", ".yml", ".ini", ".cfg", ".conf", ".toml",
    ".py", ".js", ".css", ".svg",
}
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp", ".ico"}
_PDF_EXTENSIONS = {".pdf"}
_SPREADSHEET_EXTENSIONS = {".xlsx", ".xls", ".xlsm"}
_CSV_EXTENSIONS = {".csv", ".tsv"}
_DOCX_EXTENSIONS = {".docx", ".docm"}


def classify_mimetype(mimetype: str, filename: str) -> str:
    mimetype = (mimetype or "").lower()
    ext = ""
    if filename and "." in filename:
        ext = ("." + filename.rsplit(".", 1)[-1]).lower()

    if mimetype == "application/pdf" or ext in _PDF_EXTENSIONS:
        return "pdf"
    if mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or ext in _DOCX_EXTENSIONS:
        return "docx"
    if ext in _SPREADSHEET_EXTENSIONS or mimetype in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return "spreadsheet"
    if ext in _CSV_EXTENSIONS or mimetype == "text/csv":
        return "csv"
    if mimetype.startswith("text/") or ext in _TEXT_EXTENSIONS:
        return "text"
    if mimetype.startswith("image/") or ext in _IMAGE_EXTENSIONS:
        return "image"
    return "binary"


def is_text_extractable(file_type: str) -> bool:
    return file_type in ("pdf", "docx", "spreadsheet", "csv", "text")


# ── Text extraction helpers ───────────────────────────────────────────

def extract_pdf_text(data: bytes, max_chars: int) -> tuple[str, bool]:
    import pdfplumber

    parts = []
    total = 0
    truncated = False

    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if not page_text:
                    continue
                remaining = max_chars - total
                if remaining <= 0:
                    truncated = True
                    break
                if len(page_text) > remaining:
                    parts.append(page_text[:remaining])
                    truncated = True
                    break
                parts.append(page_text)
                total += len(page_text)
    except Exception as e:
        logger.warning("pdfplumber extraction failed: %s", e)
        return "", False

    return "\n\n".join(parts), truncated


def extract_docx_text(data: bytes, max_chars: int) -> tuple[str, bool]:
    import docx

    parts: list[str] = []
    total = 0
    truncated = False

    try:
        doc = docx.Document(io.BytesIO(data))
        for para in doc.paragraphs:
            text = para.text
            if not text:
                continue
            remaining = max_chars - total
            if remaining <= 0:
                truncated = True
                break
            if len(text) > remaining:
                parts.append(text[:remaining])
                truncated = True
                break
            parts.append(text)
            total += len(text) + 1

        if not truncated:
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    line = "\t".join(cells)
                    remaining = max_chars - total
                    if remaining <= 0:
                        truncated = True
                        break
                    if len(line) > remaining:
                        parts.append(line[:remaining])
                        truncated = True
                        break
                    parts.append(line)
                    total += len(line) + 1
                if truncated:
                    break
    except Exception as e:
        logger.warning("python-docx extraction failed: %s", e)
        return "", False

    return "\n".join(parts), truncated


def extract_spreadsheet_text(data: bytes, max_chars: int) -> tuple[str, bool]:
    if data[:4] == b'\xd0\xcf\x11\xe0':
        return _extract_xls_text(data, max_chars)
    return _extract_xlsx_text(data, max_chars)


def _extract_xls_text(data: bytes, max_chars: int) -> tuple[str, bool]:
    import xlrd

    parts = []
    total = 0
    truncated = False

    try:
        wb = xlrd.open_workbook(file_contents=data)
        for sheet_name in wb.sheet_names():
            ws = wb.sheet_by_name(sheet_name)
            if wb.nsheets > 1:
                header = f"--- Sheet: {sheet_name} ---"
                parts.append(header)
                total += len(header)
            for row_idx in range(ws.nrows):
                cells = [str(c.value) if c.value != "" else "" for c in ws.row(row_idx)]
                line = "\t".join(cells)
                remaining = max_chars - total
                if remaining <= 0:
                    truncated = True
                    break
                if len(line) > remaining:
                    parts.append(line[:remaining])
                    truncated = True
                    break
                parts.append(line)
                total += len(line) + 1
            if truncated:
                break
    except Exception as e:
        logger.warning("xlrd extraction failed: %s", e)
        return f"[Spreadsheet extraction error: {e}]", False

    return "\n".join(parts), truncated


def _extract_xlsx_text(data: bytes, max_chars: int) -> tuple[str, bool]:
    from openpyxl import load_workbook

    parts = []
    total = 0
    truncated = False

    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            if len(wb.sheetnames) > 1:
                header = f"--- Sheet: {sheet_name} ---"
                parts.append(header)
                total += len(header)
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                line = "\t".join(cells)
                remaining = max_chars - total
                if remaining <= 0:
                    truncated = True
                    break
                if len(line) > remaining:
                    parts.append(line[:remaining])
                    truncated = True
                    break
                parts.append(line)
                total += len(line) + 1
            if truncated:
                break
        wb.close()
    except Exception as e:
        logger.warning("openpyxl extraction failed: %s", e)
        return f"[Spreadsheet extraction error: {e}]", False

    return "\n".join(parts), truncated


def _extract_plain_bytes(data: bytes, max_chars: int) -> tuple[str, bool]:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1")
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars], True


extract_csv_text = _extract_plain_bytes
extract_text_content = _extract_plain_bytes


def extract_text(data: bytes, file_type: str, max_chars: int) -> tuple[str, bool]:
    if file_type == "pdf":
        return extract_pdf_text(data, max_chars)
    if file_type == "docx":
        return extract_docx_text(data, max_chars)
    if file_type == "spreadsheet":
        return extract_spreadsheet_text(data, max_chars)
    if file_type == "csv":
        return extract_csv_text(data, max_chars)
    if file_type == "text":
        return extract_text_content(data, max_chars)
    return "", False


# ── Odoo helpers ──────────────────────────────────────────────────────

def decode_datas(datas) -> bytes:
    if isinstance(datas, xmlrpc.client.Binary):
        return datas.data
    if isinstance(datas, (bytes, str)):
        try:
            return base64.b64decode(datas)
        except Exception as e:
            logger.warning("Failed to base64-decode Odoo attachment datas: %s", e)
            return b""
    return b""


# ── Odoo report name mapping ─────────────────────────────────────────

_REPORT_MAP: dict[str, dict] = {
    "purchase_order": {
        "report_name": "purchase.report_purchaseorder",
        "model": "purchase.order",
        "name_field": "name",
    },
    "rfq": {
        "report_name": "purchase.report_purchasequotation",
        "model": "purchase.order",
        "name_field": "name",
    },
    "invoice": {
        "report_name": "account.report_invoice",
        "model": "account.move",
        "name_field": "name",
    },
    "sale_order": {
        "report_name": "sale.report_saleorder",
        "model": "sale.order",
        "name_field": "name",
    },
    "delivery": {
        "report_name": "stock.report_deliveryslip",
        "model": "stock.picking",
        "name_field": "name",
    },
    "picking": {
        "report_name": "stock.report_picking",
        "model": "stock.picking",
        "name_field": "name",
    },
}


# ── Tool handlers ────────────────────────────────────────────────────

def download_odoo_pdf(document_type: str, record_id: int, cache_dir: str | None = None) -> dict:
    client, err = safe_get_client()
    if not client:
        return err

    mapping = _REPORT_MAP.get(document_type)
    if not mapping:
        supported = ", ".join(sorted(_REPORT_MAP.keys()))
        return {"error": f"Unknown document type '{document_type}'. Supported: {supported}"}

    report_name = mapping["report_name"]
    model = mapping["model"]
    name_field = mapping["name_field"]

    reports = client.search_read(
        "ir.actions.report",
        [["report_name", "=", report_name]],
        ["id"],
        limit=1,
    )
    if not reports:
        return {"error": f"Report '{report_name}' not found in Odoo. The module may not be installed."}
    report_id = reports[0]["id"]

    records = client.search_read(
        model, [["id", "=", record_id]], ["id", name_field], limit=1,
    )
    if not records:
        return {"error": f"{document_type} with ID {record_id} not found"}

    display_name = records[0].get(name_field, str(record_id))
    safe_name = "".join(c if c.isalnum() or c in "-_" else "_" for c in str(display_name))
    safe_name = safe_name.strip("_") or f"{document_type}_{record_id}"
    filename = f"{safe_name}.pdf"

    try:
        result = client.execute(
            "ir.actions.report", "_render_qweb_pdf", report_id, [record_id],
        )
    except Exception as e:
        logger.error("Odoo PDF render failed for %s/%s: %s", document_type, record_id, e)
        return {"error": f"Failed to render PDF: {e}"}

    if not result:
        return {"error": "Odoo returned empty result from PDF render"}

    pdf_data = result[0] if isinstance(result, (list, tuple)) else result

    if isinstance(pdf_data, xmlrpc.client.Binary):
        raw_bytes = pdf_data.data
    elif isinstance(pdf_data, bytes):
        if pdf_data[:5] == b"%PDF-":
            raw_bytes = pdf_data
        else:
            try:
                raw_bytes = base64.b64decode(pdf_data)
            except Exception:
                raw_bytes = pdf_data
    elif isinstance(pdf_data, str):
        try:
            raw_bytes = base64.b64decode(pdf_data)
        except Exception:
            return {"error": "Failed to decode PDF data from Odoo"}
    else:
        return {"error": f"Unexpected PDF data type from Odoo: {type(pdf_data).__name__}"}

    text, truncated = extract_pdf_text(raw_bytes, 50_000)

    file_ref = None
    if cache_dir:
        from core.agents.tools.file_cache import cache_file
        file_ref = cache_file(cache_dir, raw_bytes, filename, "application/pdf")

    response: dict = {
        "ok": True,
        "document_type": document_type,
        "record_id": record_id,
        "display_name": display_name,
        "filename": filename,
        "size_bytes": len(raw_bytes),
        "text_content": text if text else None,
        "chars_extracted": len(text),
        "truncated": truncated,
    }
    if file_ref:
        response["file_ref"] = file_ref

    if not text:
        response["note"] = "PDF has no extractable text (may be a scanned/image-based document). Use file_ref to forward the original."
    elif truncated:
        response["note"] = "Content truncated at 50,000 characters"

    return response


def read_odoo_attachment(attachment_id: int, max_chars: int = 50_000) -> dict:
    client, err = safe_get_client()
    if not client:
        return err

    records = client.search_read(
        "ir.attachment",
        [["id", "=", attachment_id]],
        ["id", "name", "mimetype", "file_size", "datas"],
        limit=1,
    ) or []

    if not records:
        return {"error": f"Attachment #{attachment_id} not found"}

    rec = records[0]
    filename = rec.get("name", "")
    mimetype = rec.get("mimetype", "")
    file_size = rec.get("file_size", 0) or 0
    file_type = classify_mimetype(mimetype, filename)

    if file_size > MAX_FILE_SIZE:
        return {
            "id": attachment_id, "filename": filename, "mimetype": mimetype,
            "file_size": file_size, "file_type": file_type, "text_content": None,
            "note": f"File too large ({file_size / (1024*1024):.1f} MB, limit {MAX_FILE_SIZE / (1024*1024):.0f} MB)",
        }

    if not is_text_extractable(file_type):
        return {
            "id": attachment_id, "filename": filename, "mimetype": mimetype,
            "file_size": file_size, "file_type": file_type, "text_content": None,
            "note": f"Binary file — text extraction not supported for {mimetype or file_type}",
        }

    datas = rec.get("datas")
    if not datas:
        return {
            "id": attachment_id, "filename": filename, "mimetype": mimetype,
            "file_size": file_size, "file_type": file_type, "text_content": None,
            "note": "Attachment has no binary content",
        }

    raw = decode_datas(datas)
    if not raw:
        return {
            "id": attachment_id, "filename": filename, "mimetype": mimetype,
            "file_size": file_size, "file_type": file_type, "text_content": None,
            "note": "Failed to decode attachment data",
        }

    if len(raw) > MAX_FILE_SIZE:
        return {
            "id": attachment_id, "filename": filename, "mimetype": mimetype,
            "file_size": len(raw), "file_type": file_type, "text_content": None,
            "note": f"File too large ({len(raw) / (1024*1024):.1f} MB)",
        }

    text, truncated = extract_text(raw, file_type, max_chars)

    result = {
        "id": attachment_id, "filename": filename, "mimetype": mimetype,
        "file_size": file_size, "file_type": file_type,
        "text_content": text if text else None,
        "chars_extracted": len(text), "truncated": truncated,
    }
    if not text and file_type == "pdf":
        result["note"] = "PDF has no extractable text (may be scanned/image-based)"
    if truncated:
        result["note"] = f"Content truncated at {max_chars} characters"
    return result


def create_odoo_attachment(
    model: str,
    record_id: int,
    filename: str,
    file_ref: str = "",
    cache_dir: str | None = None,
) -> dict:
    """Create an ir.attachment on an Odoo record from a cached file."""
    from core.agents.tools.file_cache import load_cached_file

    if model not in ALLOWED_MODELS:
        return {"error": f"Model '{model}' is not allowed"}
    if not file_ref:
        return {"error": "file_ref is required — download the file first"}

    cached = load_cached_file(cache_dir, file_ref)
    if not cached:
        return {"error": f"File ref '{file_ref}' not found or expired. Download the file again."}

    raw = cached["raw"]
    mime_type = cached.get("mime_type", "application/octet-stream")

    if len(raw) > MAX_FILE_SIZE:
        return {"error": f"File too large ({len(raw)} bytes, max {MAX_FILE_SIZE})"}

    client, err = safe_get_client()
    if not client:
        return err

    datas_b64 = base64.b64encode(raw).decode("ascii")

    try:
        att_id = client.create("ir.attachment", {
            "name": filename,
            "type": "binary",
            "datas": datas_b64,
            "res_model": model,
            "res_id": record_id,
            "mimetype": mime_type,
        })
    except Exception as e:
        logger.exception("Failed to create ir.attachment on %s %d", model, record_id)
        return {"error": f"Odoo error creating attachment: {e}"}

    return {
        "ok": True,
        "attachment_id": att_id,
        "filename": filename,
        "model": model,
        "record_id": record_id,
    }


def list_record_attachments(
    model: str,
    record_id: int,
    limit: int = 20,
) -> dict:
    """List attachments on an Odoo record."""
    if model not in ALLOWED_MODELS:
        return {"error": f"Model '{model}' is not allowed"}

    client, err = safe_get_client()
    if not client:
        return err

    try:
        attachments = client.search_read(
            "ir.attachment",
            [["res_model", "=", model], ["res_id", "=", record_id]],
            ["id", "name", "mimetype", "file_size", "create_date"],
            limit=limit,
        ) or []
    except Exception as e:
        logger.exception("Failed to list attachments on %s %d", model, record_id)
        return {"error": f"Odoo error listing attachments: {e}"}

    return {
        "ok": True,
        "model": model,
        "record_id": record_id,
        "attachments": [
            {
                "id": a["id"],
                "name": a.get("name", ""),
                "mimetype": a.get("mimetype", ""),
                "file_size": a.get("file_size", 0),
                "create_date": a.get("create_date", ""),
            }
            for a in attachments
        ],
        "count": len(attachments),
    }


# ── Tool definitions ─────────────────────────────────────────────────

PDF_TOOL_DEFS = [
    {
        "name": "download_odoo_pdf",
        "description": (
            "Download a PDF copy of an Odoo document (purchase order, RFQ, invoice, sale order, "
            "delivery slip, or picking). Returns extracted text content and a file_ref for "
            "forwarding the original PDF via send/reply email tools."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "document_type": {
                    "type": "string",
                    "enum": ["purchase_order", "rfq", "invoice", "sale_order", "delivery", "picking"],
                    "description": "Type of Odoo document to download as PDF. Use 'rfq' for draft/sent purchase orders, 'purchase_order' for confirmed POs.",
                },
                "record_id": {
                    "type": "integer",
                    "description": "Odoo record ID of the document",
                },
            },
            "required": ["document_type", "record_id"],
        },
        "kind": "integration",
        "writes": False,
    },
    {
        "name": "read_odoo_attachment",
        "description": (
            "Read and extract text from an Odoo attachment by its ID. Supports PDFs, Word docs "
            "(.docx), spreadsheets (.xlsx/.xls), CSV, and text files. Use this to read documents "
            "attached to POs, invoices, or other Odoo records."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "attachment_id": {
                    "type": "integer",
                    "description": "Odoo ir.attachment record ID",
                },
                "max_chars": {
                    "type": "integer",
                    "description": "Maximum characters to extract (default 50000)",
                },
            },
            "required": ["attachment_id"],
        },
        "kind": "integration",
        "writes": False,
    },
    {
        "name": "create_odoo_attachment",
        "description": (
            "Create an attachment on an Odoo record from a previously downloaded file. "
            "Pass the file_ref from download_email_attachment or download_odoo_pdf. "
            "Use this to attach COAs, vendor documents, or other files to POs, invoices, etc."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "model": {
                    "type": "string",
                    "description": "Odoo model to attach to (e.g. 'purchase.order')",
                },
                "record_id": {
                    "type": "integer",
                    "description": "Record ID to attach the file to",
                },
                "filename": {
                    "type": "string",
                    "description": "Filename for the attachment (e.g. 'COA-12345.pdf')",
                },
                "file_ref": {
                    "type": "string",
                    "description": "File reference from download_email_attachment or download_odoo_pdf",
                },
            },
            "required": ["model", "record_id", "filename", "file_ref"],
        },
        "kind": "integration",
        "writes": True,
    },
    {
        "name": "list_record_attachments",
        "description": (
            "List attachments on an Odoo record. Returns file names, IDs, and sizes. "
            "Use read_odoo_attachment to read the content of a specific attachment."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "model": {
                    "type": "string",
                    "description": "Odoo model name (e.g. 'purchase.order')",
                },
                "record_id": {
                    "type": "integer",
                    "description": "Record ID to list attachments for",
                },
                "limit": {
                    "type": "integer",
                    "description": "Max attachments to return (default 20)",
                },
            },
            "required": ["model", "record_id"],
        },
        "kind": "integration",
        "writes": False,
    },
]

PDF_EXECUTORS = {
    "download_odoo_pdf": download_odoo_pdf,
    "read_odoo_attachment": read_odoo_attachment,
    "create_odoo_attachment": create_odoo_attachment,
    "list_record_attachments": list_record_attachments,
}
