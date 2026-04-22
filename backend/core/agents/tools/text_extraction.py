"""Shared text extraction utilities.

File-type classification and text extraction for PDFs, Word docs (.docx),
spreadsheets (.xlsx/.xls), CSV, and plain text. Used by both Gmail
attachment tools and Odoo attachment tools.
"""

import io
import logging

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB

# ── File-type classification ──────────────────────────────────────────

_TEXT_EXTENSIONS = {
    ".txt", ".eml", ".html", ".htm", ".xml", ".json", ".log",
    ".md", ".yaml", ".yml", ".ini", ".cfg", ".conf", ".toml",
    ".py", ".js", ".css", ".svg",
}
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp", ".ico"}
_PDF_EXTENSIONS = {".pdf"}
_SPREADSHEET_EXTENSIONS = {".xlsx", ".xls", ".xlsm"}
_CSV_EXTENSIONS = {".csv", ".tsv"}
_DOCX_EXTENSIONS = {".docx", ".docm"}


def classify_mimetype(mimetype: str, filename: str) -> str:
    mimetype = (mimetype or "").lower()
    ext = ""
    if filename and "." in filename:
        ext = ("." + filename.rsplit(".", 1)[-1]).lower()

    if mimetype == "application/pdf" or ext in _PDF_EXTENSIONS:
        return "pdf"
    if mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or ext in _DOCX_EXTENSIONS:
        return "docx"
    if ext in _SPREADSHEET_EXTENSIONS or mimetype in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return "spreadsheet"
    if ext in _CSV_EXTENSIONS or mimetype == "text/csv":
        return "csv"
    if mimetype.startswith("text/") or ext in _TEXT_EXTENSIONS:
        return "text"
    if mimetype.startswith("image/") or ext in _IMAGE_EXTENSIONS:
        return "image"
    return "binary"


def is_text_extractable(file_type: str) -> bool:
    return file_type in ("pdf", "docx", "spreadsheet", "csv", "text")


# ── Text extraction helpers ───────────────────────────────────────────

def extract_pdf_text(data: bytes, max_chars: int) -> tuple[str, bool]:
    import pdfplumber

    parts = []
    total = 0
    truncated = False

    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if not page_text:
                    continue
                remaining = max_chars - total
                if remaining <= 0:
                    truncated = True
                    break
                if len(page_text) > remaining:
                    parts.append(page_text[:remaining])
                    truncated = True
                    break
                parts.append(page_text)
                total += len(page_text)
    except Exception as e:
        logger.warning("pdfplumber extraction failed: %s", e)
        return "", False

    return "\n\n".join(parts), truncated


def extract_docx_text(data: bytes, max_chars: int) -> tuple[str, bool]:
    import docx

    parts: list[str] = []
    total = 0
    truncated = False

    try:
        doc = docx.Document(io.BytesIO(data))
        for para in doc.paragraphs:
            text = para.text
            if not text:
                continue
            remaining = max_chars - total
            if remaining <= 0:
                truncated = True
                break
            if len(text) > remaining:
                parts.append(text[:remaining])
                truncated = True
                break
            parts.append(text)
            total += len(text) + 1

        if not truncated:
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    line = "\t".join(cells)
                    remaining = max_chars - total
                    if remaining <= 0:
                        truncated = True
                        break
                    if len(line) > remaining:
                        parts.append(line[:remaining])
                        truncated = True
                        break
                    parts.append(line)
                    total += len(line) + 1
                if truncated:
                    break
    except Exception as e:
        logger.warning("python-docx extraction failed: %s", e)
        return "", False

    return "\n".join(parts), truncated


def extract_spreadsheet_text(data: bytes, max_chars: int) -> tuple[str, bool]:
    if data[:4] == b'\xd0\xcf\x11\xe0':
        return _extract_xls_text(data, max_chars)
    return _extract_xlsx_text(data, max_chars)


def _extract_xls_text(data: bytes, max_chars: int) -> tuple[str, bool]:
    import xlrd

    parts = []
    total = 0
    truncated = False

    try:
        wb = xlrd.open_workbook(file_contents=data)
        for sheet_name in wb.sheet_names():
            ws = wb.sheet_by_name(sheet_name)
            if wb.nsheets > 1:
                header = f"--- Sheet: {sheet_name} ---"
                parts.append(header)
                total += len(header)
            for row_idx in range(ws.nrows):
                cells = [str(c.value) if c.value != "" else "" for c in ws.row(row_idx)]
                line = "\t".join(cells)
                remaining = max_chars - total
                if remaining <= 0:
                    truncated = True
                    break
                if len(line) > remaining:
                    parts.append(line[:remaining])
                    truncated = True
                    break
                parts.append(line)
                total += len(line) + 1
            if truncated:
                break
    except Exception as e:
        logger.warning("xlrd extraction failed: %s", e)
        return f"[Spreadsheet extraction error: {e}]", False

    return "\n".join(parts), truncated


def _extract_xlsx_text(data: bytes, max_chars: int) -> tuple[str, bool]:
    from openpyxl import load_workbook

    parts = []
    total = 0
    truncated = False

    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            if len(wb.sheetnames) > 1:
                header = f"--- Sheet: {sheet_name} ---"
                parts.append(header)
                total += len(header)
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                line = "\t".join(cells)
                remaining = max_chars - total
                if remaining <= 0:
                    truncated = True
                    break
                if len(line) > remaining:
                    parts.append(line[:remaining])
                    truncated = True
                    break
                parts.append(line)
                total += len(line) + 1
            if truncated:
                break
        wb.close()
    except Exception as e:
        logger.warning("openpyxl extraction failed: %s", e)
        return f"[Spreadsheet extraction error: {e}]", False

    return "\n".join(parts), truncated


def _extract_plain_bytes(data: bytes, max_chars: int) -> tuple[str, bool]:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1")
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars], True


extract_csv_text = _extract_plain_bytes
extract_text_content = _extract_plain_bytes


def extract_text(data: bytes, file_type: str, max_chars: int) -> tuple[str, bool]:
    if file_type == "pdf":
        return extract_pdf_text(data, max_chars)
    if file_type == "docx":
        return extract_docx_text(data, max_chars)
    if file_type == "spreadsheet":
        return extract_spreadsheet_text(data, max_chars)
    if file_type == "csv":
        return extract_csv_text(data, max_chars)
    if file_type == "text":
        return extract_text_content(data, max_chars)
    return "", False
